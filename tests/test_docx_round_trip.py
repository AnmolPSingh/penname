"""DOCX path: paragraph- and table-aware pseudonymization, text-level round trip."""

from pathlib import Path

import pytest
from docx import Document

from penname.core.engine import PennameSession
from penname.core.io.docx_io import pseudonymize_docx
from penname.core.replace.applier import reverse_text


@pytest.fixture
def document_path(tmp_path: Path) -> Path:
    doc = Document()
    doc.add_heading("Stewardship Notes", level=1)
    p = doc.add_paragraph("Met with ")
    p.add_run("Robert Castellano").bold = True  # entity split across styled runs
    p.add_run(" at his home in Santa Fe on June 3, 2025.")
    doc.add_paragraph(
        "Follow up with his daughter Maria Castellano-Reyes at (505) 555-0177."
    )
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Donor"
    table.cell(0, 1).text = "Email"
    table.cell(1, 0).text = "Priya Raghunathan"
    table.cell(1, 1).text = "priya.r@fastmail.com"

    path = tmp_path / "notes.docx"
    doc.save(path)
    return path


def _all_texts(doc: Document) -> list[str]:
    texts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.extend(p.text for p in cell.paragraphs)
    return texts


def test_docx_text_round_trips(document_path: Path, tmp_path: Path) -> None:
    out = tmp_path / "notes.pen.docx"
    mapping = pseudonymize_docx(document_path, out, PennameSession())

    original_texts = _all_texts(Document(str(document_path)))
    pen_texts = _all_texts(Document(str(out)))
    assert len(pen_texts) == len(original_texts)

    for original, pen in zip(original_texts, pen_texts):
        assert reverse_text(pen, mapping) == original


def test_docx_headers_and_footers_are_pseudonymized(tmp_path: Path) -> None:
    """Donor letterhead and reply-to lines live in headers/footers — they must
    be pseudonymized, not silently passed through."""
    doc = Document()
    doc.sections[0].header.paragraphs[0].text = "Riverside Foundation — Margaret Wilson"
    doc.add_paragraph("Thank you for your generous gift.")
    doc.sections[0].footer.paragraphs[0].text = "Reply to s.chen@riversidecf.org"
    source = tmp_path / "letter.docx"
    doc.save(source)
    out = tmp_path / "letter.pen.docx"

    mapping = pseudonymize_docx(source, out, PennameSession())

    pen = Document(str(out))
    header = pen.sections[0].header.paragraphs[0].text
    footer = pen.sections[0].footer.paragraphs[0].text
    assert "Margaret Wilson" not in header
    assert "s.chen@riversidecf.org" not in footer
    # and both reverse cleanly
    assert reverse_text(header, mapping) == "Riverside Foundation — Margaret Wilson"
    assert reverse_text(footer, mapping) == "Reply to s.chen@riversidecf.org"


def test_docx_structure_preserved_and_names_replaced(
    document_path: Path, tmp_path: Path
) -> None:
    out = tmp_path / "notes.pen.docx"
    pseudonymize_docx(document_path, out, PennameSession())

    original = Document(str(document_path))
    pen = Document(str(out))
    assert len(pen.paragraphs) == len(original.paragraphs)
    assert len(pen.tables) == len(original.tables)

    joined = "\n".join(_all_texts(pen))
    assert "Robert Castellano" not in joined  # even though split across runs
    assert "priya.r@fastmail.com" not in joined
