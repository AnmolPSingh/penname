"""Markdown export — the LLM-friendly copy for every format."""

from pathlib import Path

from penname.core.io.markdown import export_markdown, to_markdown



def _utf8(text: str) -> bytes:
    """write_text() would translate \\n to the platform separator; these tests
    assert on exact bytes, so they must write exact bytes."""
    return text.encode("utf-8")


def test_csv_becomes_a_table(tmp_path: Path) -> None:
    source = tmp_path / "donors.csv"
    source.write_bytes(_utf8("Name,City\nDorothy Fields,Springfield\n"))

    md = to_markdown(source)

    assert "| Name | City |" in md
    assert "| --- | --- |" in md
    assert "| Dorothy Fields | Springfield |" in md


def test_xlsx_becomes_tables_per_sheet(tmp_path: Path) -> None:
    from openpyxl import Workbook

    wb = Workbook()
    wb.active.title = "Donors"
    wb.active.append(["Name", "Gift"])
    wb.active.append(["Dorothy Fields", 100])
    source = tmp_path / "donors.xlsx"
    wb.save(source)

    md = to_markdown(source)

    assert "## Donors" in md
    assert "| Dorothy Fields | 100 |" in md


def test_docx_becomes_paragraphs_and_tables(tmp_path: Path) -> None:
    from docx import Document

    doc = Document()
    doc.add_paragraph("Met with Dorothy Fields.")
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Donor"
    table.cell(0, 1).text = "Dorothy Fields"
    source = tmp_path / "notes.docx"
    doc.save(source)

    md = to_markdown(source)

    assert "Met with Dorothy Fields." in md
    assert "| Donor | Dorothy Fields |" in md


def test_text_passes_through_and_pipes_are_escaped(tmp_path: Path) -> None:
    source = tmp_path / "notes.txt"
    source.write_bytes(_utf8("Plain text stays as is.\n"))
    assert to_markdown(source) == "Plain text stays as is.\n"

    csv_source = tmp_path / "weird.csv"
    csv_source.write_bytes(_utf8('Notes\n"has | a pipe"\n'))
    assert "has \\| a pipe" in to_markdown(csv_source)


def test_export_markdown_writes_file(tmp_path: Path) -> None:
    source = tmp_path / "notes.md"
    source.write_bytes(_utf8("# Heading\n"))
    dest = tmp_path / "out.md"

    export_markdown(source, dest)

    assert dest.read_text(encoding="utf-8") == "# Heading\n"
