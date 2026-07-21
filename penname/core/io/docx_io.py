"""DOCX path: paragraph-level pseudonymization, run-aware.

Detected values can cross styled runs, so each changed paragraph is rewritten
as a single run carrying the paragraph's first-run formatting. Unchanged
paragraphs keep their runs untouched.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph

from penname.core.engine import PennameSession, RoundTripError
from penname.core.replace.applier import reverse_text
from penname.core.types import Mapping, MappingEntry


def _paragraphs_in(container):
    """Every paragraph in a body/header/footer, including those inside tables."""
    yield from container.paragraphs
    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def _iter_paragraphs(doc: Document):
    """All paragraphs including headers and footers — donor letterhead and
    reply-to lines live there and must not be skipped. Headers/footers that are
    linked to the previous section share the same underlying part, so they are
    skipped to avoid processing (and re-pseudonymizing) the same text twice."""
    yield from _paragraphs_in(doc)
    for section in doc.sections:
        for hdrftr in (
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
            section.even_page_header,
            section.even_page_footer,
        ):
            if not hdrftr.is_linked_to_previous:
                yield from _paragraphs_in(hdrftr)


def _rewrite(paragraph: Paragraph, new_text: str) -> None:
    first_run_style = paragraph.runs[0].style if paragraph.runs else None
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)
    run = paragraph.add_run(new_text)
    if first_run_style is not None:
        run.style = first_run_style


def pseudonymize_docx(
    source: str | Path, dest: str | Path, session: PennameSession
) -> Mapping:
    doc = Document(str(source))
    entries: dict[tuple[str, str], MappingEntry] = {}
    # Every paragraph, changed or not: untouched text must also survive the
    # union mapping's reversal (see xlsx_io for the failure mode).
    processed: list[tuple[str, str]] = []

    for paragraph in _iter_paragraphs(doc):
        text = paragraph.text
        if not text.strip():
            continue
        result = session.pseudonymize(text)
        for entry in result.mapping.entries:
            entries[(entry.entity_type, entry.original)] = entry
        processed.append((text, result.text))
        if result.text != text:
            _rewrite(paragraph, result.text)

    mapping = Mapping(entries=tuple(entries.values()))
    for original, pseudonymized in processed:
        if reverse_text(pseudonymized, mapping) != original:
            raise RoundTripError(
                "a value in this document could not be pseudonymized reversibly"
            )

    doc.save(str(dest))
    return mapping
