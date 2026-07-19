"""Markdown export — the LLM-friendly copy, always available.

Converts an already-pseudonymized output file; it never sees real values.
"""

from __future__ import annotations

import csv
from pathlib import Path


def _escape_cell(value: str) -> str:
    return value.replace("|", "\\|").replace("\n", " ")


def _table(rows: list[list[str]]) -> str:
    if not rows:
        return ""
    width = max(len(r) for r in rows)
    padded = [r + [""] * (width - len(r)) for r in rows]
    lines = ["| " + " | ".join(_escape_cell(c) for c in padded[0]) + " |"]
    lines.append("|" + " --- |" * width)
    lines.extend("| " + " | ".join(_escape_cell(c) for c in r) + " |" for r in padded[1:])
    return "\n".join(lines) + "\n"


def to_markdown(source: str | Path) -> str:
    source = Path(source)
    suffix = source.suffix.lower()

    if suffix == ".csv":
        with open(source, newline="", encoding="utf-8") as f:
            return _table([list(map(str, row)) for row in csv.reader(f)])

    if suffix == ".xlsx":
        from openpyxl import load_workbook

        parts = []
        for ws in load_workbook(source, read_only=True).worksheets:
            rows = [
                ["" if c is None else str(c) for c in row]
                for row in ws.iter_rows(values_only=True)
            ]
            parts.append(f"## {ws.title}\n\n{_table(rows)}")
        return "\n".join(parts)

    if suffix == ".docx":
        from docx import Document

        doc = Document(str(source))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        parts.extend(
            _table([[cell.text for cell in row.cells] for row in t.rows])
            for t in doc.tables
        )
        return "\n\n".join(parts) + "\n"

    return source.read_bytes().decode("utf-8")


def export_markdown(source: str | Path, dest: str | Path) -> None:
    Path(dest).write_bytes(to_markdown(source).encode("utf-8"))
